
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from datetime import date
class Template:
    def __init__(self,title):
        # Create Document
        self.doc = Document()

        # Create header
        header = self.doc.sections[0].header
        header_paragraph = header.paragraphs[0]
        header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        current_date = date.today().strftime("%B %d, %Y")
        date_run = header_paragraph.add_run(current_date)
        date_run.font.size = Pt(12)

        # Create title
        self.title = title
        subh = input('Subheading?')
        author = input('Who is it by?')
        self.doc.add_heading(self.title, level=0)
        self.doc.add_heading(subh, level=1)

        # Create footer
        footer = self.doc.sections[0].footer
        footer_paragraph = footer.paragraphs[0]
        footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        footerr = footer_paragraph.add_run('By: ' + author)
        footerr.font.size = Pt(12)

        self.doc.add_page_break()
        self.doc.save(self.title + '.docx')

    def add_headings(self):
        sections = []
        while True:
            v = input("Enter section titles or 'done' to finish: ")
            if v.lower() == "done":
                self.doc.save(self.title + '.docx')
                break
            sections.append(v)
            self.doc.add_heading(v + ':',level=3)
