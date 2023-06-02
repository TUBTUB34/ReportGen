
from docx import Document
from docx.enum.style import WD_STYLE_TYPE

class Edit:
    def __init__(self,Title):
        self.title = Title
        self.doc = Document(Title)
        self.headings = []

        for paragraph in self.doc.paragraphs:
            if paragraph.style.name == 'Heading 3':
                self.headings.append(paragraph.text)


    def add_toc(self):
    

        # Save the document
        self.doc.save('your_document_with_toc.docx')

    def add_info(self,horp):
        self.print_headings()
        while True:
            if (horp == 'p'):
                temp = input("Enter the heading where you want to add a new paragraph(done to return): ")
            if (horp == 'h'):
                temp = input("Enter the heading you want to add the heading below (done to return): ")
            target_heading = temp+':'
            
            if temp == 'done':
                break

            if target_heading not in self.headings:
                print("Heading not found. Please try again.")
                continue

            if (horp == 'p'):
                new_paragraph_text = input("Enter the text for the new paragraph: ")
            if (horp == 'h'):
                new_paragraph_text = input("Enter the text for the new heading: ")

            paragraph_index = 0

            for i, paragraph in enumerate(self.doc.paragraphs):
                if paragraph.text.lower() == target_heading.lower():
                    paragraph_index = i
                    break

            next_heading_index = len(self.doc.paragraphs)
            for i in range(paragraph_index + 1, len(self.doc.paragraphs)):
                if self.doc.paragraphs[i].text in self.headings:
                    next_heading_index = i
                    break
            if (horp == 'p'):
                self.doc.paragraphs[next_heading_index].insert_paragraph_before(new_paragraph_text)
                print("New paragraph added successfully.")
            if (horp == 'h'):
                h2 = self.doc.paragraphs[next_heading_index].insert_paragraph_before(new_paragraph_text)
                h2.style = self.doc.styles['Heading 3']
                print("heading added successfully.")
            
            break

        self.doc.save('modified_document.docx')
    

    def save_document(self, filename):
        self.doc.save(filename)

    def print_headings(self):
        print('\nHeadings in the document:')
        for i in self.headings:
            print(i)